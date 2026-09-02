"""
Genera CourseTrack_EspecificacionesTecnicas.docx
Especificaciones técnicas y buenas prácticas de la aplicación CourseTrack.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colores ───────────────────────────────────────────────────────────────
AZUL       = RGBColor(0x1E, 0x3A, 0x5F)
AZUL_DARK  = RGBColor(0x15, 0x2A, 0x46)
AZUL_MID   = RGBColor(0x25, 0x63, 0xEB)
GRIS_TEXT  = RGBColor(0x47, 0x55, 0x69)
GRIS_LIGHT = RGBColor(0x94, 0xA3, 0xB8)
BLANCO     = RGBColor(0xFF, 0xFF, 0xFF)
VERDE      = RGBColor(0x16, 0xA3, 0x4A)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def make_table(doc, headers, rows, col_widths_cm, shade_rows=True):
    n = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Cm(col_widths_cm[i])
        set_cell_bg(cell, '1E3A5F')
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = BLANCO
        run.font.size = Pt(9)
        run.font.name = 'Segoe UI'

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = 'F7F9FC' if shade_rows and ri % 2 == 1 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.width = Cm(col_widths_cm[ci])
            set_cell_bg(cell, bg)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(val)
            run.font.color.rgb = GRIS_TEXT
            run.font.size = Pt(9)
            run.font.name = 'Segoe UI'

    return table


def add_heading(doc, text, level=1, color=None):
    if color is None:
        color = AZUL if level == 1 else AZUL_DARK
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.name = 'Segoe UI'
    return p


def add_para(doc, text, color=None, bold=False, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Segoe UI'
    run.font.color.rgb = color or GRIS_TEXT
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, bold_start='', level=0, color=None):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    if bold_start and text.startswith(bold_start):
        r1 = p.add_run(bold_start)
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.name = 'Segoe UI'
        r1.font.color.rgb = color or GRIS_TEXT
        r2 = p.add_run(text[len(bold_start):])
        r2.font.size = Pt(10)
        r2.font.name = 'Segoe UI'
        r2.font.color.rgb = color or GRIS_TEXT
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Segoe UI'
        run.font.color.rgb = color or GRIS_TEXT


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        # fondo gris claro
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F1F5F9')
        pPr.append(shd)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def add_spacer(doc, pts=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)


# ─────────────────────────────────────────────────────────────────────────
#  DOCUMENTO
# ─────────────────────────────────────────────────────────────────────────
doc = Document()

# Márgenes
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# Fuente por defecto
style = doc.styles['Normal']
style.font.name = 'Segoe UI'
style.font.size = Pt(10)
style.font.color.rgb = GRIS_TEXT

# ── Portada ───────────────────────────────────────────────────────────────
add_spacer(doc, 40)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run('CourseTrack')
r.bold = True; r.font.size = Pt(36); r.font.color.rgb = AZUL; r.font.name = 'Segoe UI'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(4)
r2 = p2.add_run('Especificaciones técnicas y buenas prácticas')
r2.font.size = Pt(16); r2.font.color.rgb = GRIS_TEXT; r2.font.name = 'Segoe UI'

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(30)
r3 = p3.add_run('Gestión del ciclo de vida de cursos eLearning')
r3.font.size = Pt(11); r3.font.color.rgb = GRIS_LIGHT; r3.font.name = 'Segoe UI'

# ── 1. Visión general ─────────────────────────────────────────────────────
add_heading(doc, '1. Visión general', 1)
add_para(doc,
    'CourseTrack es una aplicación web de página única (SPA) diseñada para gestionar el '
    'ciclo de vida completo de cursos eLearning: desde la recepción del material del proveedor '
    'hasta su publicación en el LMS. Permite controlar estados, versiones, incidencias y '
    'trazabilidad por usuario sin necesidad de infraestructura de servidor propia.')
add_spacer(doc, 6)
make_table(doc,
    ['Característica', 'Valor'],
    [
        ['Tipo de aplicación',        'SPA — fichero HTML único autocontenido'],
        ['Hosting',                   'GitHub Pages (rama main, raíz /)'],
        ['Backend',                   'Google Apps Script (REST)'],
        ['Persistencia local',        'localStorage del navegador'],
        ['Persistencia compartida',   'Google Sheets (JSON en celda A1)'],
        ['Frameworks JS',             'Ninguno — Vanilla JS / HTML5 / CSS3'],
        ['Dependencias externas',     'Ninguna — sin CDN, sin npm en runtime'],
        ['Soporte responsive',        'Sí — CSS Grid + Flexbox + Media Queries'],
    ],
    [6, 10],
)
add_spacer(doc, 12)

# ── 2. Tecnologías ────────────────────────────────────────────────────────
add_heading(doc, '2. Tecnologías utilizadas', 1)

add_heading(doc, '2.1 Frontend', 2)
add_para(doc,
    'La aplicación vive íntegramente en un único fichero index.html (≈ 85 KB). '
    'No hay proceso de compilación ni bundler.')
add_bullet(doc, 'HTML5 semántico — estructura, modales, tablas y formularios.')
add_bullet(doc, 'CSS3 — variables CSS (custom properties), Grid, Flexbox, Media Queries y transiciones. Sin preprocesador.')
add_bullet(doc, 'JavaScript ES2020+ Vanilla — async/await, arrow functions, destructuring, template literals. Sin frameworks ni librerías externas.')
add_bullet(doc, 'Tipografía del sistema: -apple-system, BlinkMacSystemFont, Segoe UI — sin fuentes externas.')
add_spacer(doc)

add_heading(doc, '2.2 Persistencia', 2)
add_para(doc, 'La aplicación gestiona dos capas de persistencia de forma transparente:')
add_spacer(doc, 4)
make_table(doc,
    ['Capa', 'Clave / Endpoint', 'Descripción'],
    [
        ['localStorage',   'coursetrack_v1',        'Datos del curso (JSON). Fuente de verdad local.'],
        ['localStorage',   'coursetrack_gas_url',    'URL del Web App de Google Apps Script.'],
        ['localStorage',   'coursetrack_user',       'Nombre del usuario activo en ese navegador.'],
        ['Google Sheets',  'POST/GET al GAS URL',    'JSON de todos los cursos en la celda A1 de la pestaña Data.'],
    ],
    [3.5, 4.5, 8],
)
add_spacer(doc, 6)
add_para(doc,
    'La capa local se escribe siempre de forma síncrona. La capa GAS es opcional y asíncrona: '
    'si falla, los datos permanecen en localStorage y se muestra un aviso. '
    'El modelo de concurrencia es last-write-wins — válido para equipos pequeños (< 5 personas simultáneas).')
add_spacer(doc)

add_heading(doc, '2.3 Backend — Google Apps Script', 2)
add_bullet(doc, 'Web App desplegada con acceso «Cualquiera (sin iniciar sesión)».')
add_bullet(doc, 'GET: devuelve el JSON almacenado en A1 (array de cursos).')
add_bullet(doc,
    'POST: recibe el JSON completo y lo sobreescribe en A1. '
    'Registra cada escritura en la pestaña Log con fecha/hora y tamaño en bytes.')
add_bullet(doc,
    'Comunicación con fetch() en modo no-cors (fire-and-forget) para evitar el preflight CORS '
    'cuando la app se abre desde file://.')
add_spacer(doc)

add_heading(doc, '2.4 Hosting — GitHub Pages', 2)
add_bullet(doc, 'Repositorio: github.com/pacosirvent80-hash/accenture-coursetrack')
add_bullet(doc, 'Rama de publicación: main, raíz /.')
add_bullet(doc, 'Actualización automática: 1–2 minutos tras git push.')
add_bullet(doc, 'No se necesita Dockerfile, CI/CD ni proceso de build.')
add_spacer(doc, 12)

# ── 3. Arquitectura y modelo de datos ─────────────────────────────────────
add_heading(doc, '3. Arquitectura y modelo de datos', 1)

add_heading(doc, '3.1 Estructura de datos', 2)
add_para(doc, 'Los datos se almacenan como un array JSON de objetos Course:')
add_spacer(doc, 4)
add_code_block(doc, [
    'Course {',
    '  id            : string   // UUID generado al crear',
    '  name          : string   // nombre del curso',
    '  code          : string   // código interno',
    '  provider      : string   // proveedor (de PROVIDERS[])',
    '  scormVersion  : string   // SCORM 1.2 | SCORM 2004 | xAPI | HTML',
    '  duration      : string   // duración estimada',
    '  language      : string   // idioma (de LANGUAGES[])',
    '  notes         : string   // notas generales',
    '  createdAt     : string   // fecha ISO de creación',
    '  updatedAt     : string   // fecha ISO de última modificación',
    '  versions[]    : Version',
    '}',
    '',
    'Version {',
    '  id              : string',
    '  versionNumber   : string   // "1.0", "1.1", "2.0"',
    '  status          : string   // entregado|en_revision|incidencia|...',
    '  date            : string',
    '  notes           : string',
    '  modificadoPor   : string   // usuario activo al crear la versión',
    '  incidents[]     : Incident',
    '}',
    '',
    'Incident {',
    '  id            : string',
    '  title         : string',
    '  description   : string',
    '  category      : string   // SCORM|Contenido|Diseño/UX|...',
    '  severity      : string   // Alta|Media|Baja',
    '  date          : string',
    '  resolved      : boolean',
    '  resolution    : string',
    '  creadaPor     : string',
    '  editadoPor    : string',
    '  fechaEdicion  : string',
    '  resueltaPor   : string',
    '}',
])

add_heading(doc, '3.2 Estados del ciclo de vida', 2)
add_para(doc, 'Cada versión de un curso pasa por los siguientes estados:')
add_spacer(doc, 4)
make_table(doc,
    ['Estado', 'Color', 'Descripción'],
    [
        ['entregado',      'Índigo',   'Primera entrega del proveedor.'],
        ['en_revision',    'Violeta',  'El equipo está revisando el contenido.'],
        ['incidencia',     'Rojo',     'Se han detectado errores. Hay incidencias abiertas.'],
        ['en_correccion',  'Naranja',  'El proveedor está corrigiendo los errores reportados.'],
        ['aprobado',       'Verde',    'El contenido ha superado la revisión.'],
        ['publicado',      'Teal',     'El curso está activo en el LMS.'],
        ['en_pausa',       'Gris',     'Proyecto pausado temporalmente.'],
    ],
    [4.5, 2.5, 9],
)
add_spacer(doc, 12)

# ── 4. Configuración de listas predefinidas ───────────────────────────────
add_heading(doc, '4. Configuración de listas predefinidas', 1)
add_para(doc,
    'Tres arrays en la cabecera de index.html (≈ línea 806) controlan los valores '
    'disponibles en los selectores. Son la única configuración que un no-desarrollador '
    'necesita tocar.')
add_spacer(doc, 4)
add_code_block(doc, [
    '// ── Editar estas listas para añadir/quitar valores ──',
    "const USERS = [",
    "  'Paco',",
    "  'Lourdes',",
    "  'Natalia',",
    "  // añadir más miembros del equipo",
    "];",
    "",
    "const PROVIDERS = [",
    "  'Accenture',",
    "  'INAP',",
    "  'Mercadona',",
    "  'Test Provider',",
    "];",
    "",
    "const LANGUAGES = [",
    "  'Español',",
    "  'Inglés',",
    "  'Catalán',",
    "  'Euskera',",
    "  'Gallego',",
    "];",
])
make_table(doc,
    ['Array', 'Dónde se usa', 'Comportamiento'],
    [
        ['USERS',
         'Modal de selección al primer acceso',
         'Bloqueado tras la selección. Para cambiarlo: borrar coursetrack_user en localStorage.'],
        ['PROVIDERS',
         'Select en formulario nuevo/editar curso; filtro del toolbar',
         'Cursos con proveedor fuera del array siguen siendo visibles (compatibilidad).'],
        ['LANGUAGES',
         'Select en formulario nuevo/editar curso',
         'El primer valor de la lista es el valor por defecto.'],
    ],
    [3.5, 5, 7.5],
)
add_spacer(doc, 12)

# ── 5. Seguridad ──────────────────────────────────────────────────────────
add_heading(doc, '5. Seguridad y privacidad', 1)

add_heading(doc, '5.1 URL del Apps Script', 2)
add_para(doc,
    'La URL del Web App de Google Apps Script otorga acceso de lectura y escritura '
    'sin autenticación. Nunca debe aparecer en el código fuente del repositorio público.')
add_bullet(doc, 'Cada usuario la introduce manualmente en el modal ⚙ Configuración.')
add_bullet(doc, 'Se guarda únicamente en el localStorage de ese navegador (clave coursetrack_gas_url).')
add_bullet(doc, 'No se incluye en git ni en ningún fichero del proyecto.')
add_spacer(doc)

add_heading(doc, '5.2 Modelo de autenticación', 2)
add_para(doc,
    'La selección de usuario NO es autenticación: cualquier persona con acceso a la URL '
    'puede elegir cualquier nombre. El sistema está diseñado para equipos de confianza '
    'donde la trazabilidad importa más que la seguridad de acceso.')
add_bullet(doc, 'El nombre queda registrado en todos los cambios (modificadoPor, creadaPor, editadoPor, resueltaPor).')
add_bullet(doc, 'Para cambiar el usuario: borrar coursetrack_user en DevTools → Application → Local Storage.')
add_spacer(doc)

add_heading(doc, '5.3 Concurrencia', 2)
add_bullet(doc, 'Modelo last-write-wins sobre la celda A1 de Google Sheets.')
add_bullet(doc, 'Si dos usuarios guardan simultáneamente, gana el último en llegar.')
add_bullet(doc, 'La pestaña Log de Google Sheets registra cada escritura con fecha/hora para auditoría.')
add_spacer(doc, 12)

# ── 6. Buenas prácticas ───────────────────────────────────────────────────
add_heading(doc, '6. Buenas prácticas de trabajo', 1)

add_heading(doc, '6.1 Modificar el código fuente', 2)
add_bullet(doc, 'Trabajar siempre sobre index.html en el repositorio local.')
add_bullet(doc, 'Usar VS Code con la extensión Prettier para no romper el formato.')
add_bullet(doc, 'No externalizar CSS o JS — la autocontención es un requisito de diseño.')
add_bullet(doc, 'Probar en el navegador local (abrir index.html) antes de hacer push.')
add_bullet(doc, 'Hacer git push solo cuando el cambio está verificado — GitHub Pages publica en 1–2 minutos.')
add_spacer(doc)

add_heading(doc, '6.2 Añadir o quitar usuarios / proveedores / idiomas', 2)
add_bullet(doc, 'Abrir index.html en el editor.')
add_bullet(doc, 'Localizar el array correspondiente (≈ línea 806): USERS, PROVIDERS o LANGUAGES.')
add_bullet(doc, 'Añadir o eliminar la cadena de texto deseada.')
add_bullet(doc, 'Guardar → git add index.html → git commit → git push.')
add_para(doc,
    'Nota: eliminar un proveedor del array no borra los cursos que ya lo usan. '
    'Seguirán apareciendo en los filtros mientras existan en los datos.',
    color=GRIS_LIGHT, space_after=6)
add_spacer(doc)

add_heading(doc, '6.3 Gestionar la URL del Apps Script', 2)
add_bullet(doc, 'Configurar en el modal ⚙ (botón en la cabecera). Se guarda en el localStorage de ese navegador.')
add_bullet(doc, 'Cada miembro del equipo introduce la URL en su propio navegador — nunca compartirla por canales públicos.')
add_bullet(doc, 'Si se redespliega el Apps Script, la URL cambia: actualizar en todos los navegadores del equipo.')
add_bullet(doc, 'Para desconectar de Sheets: dejar el campo URL vacío en ⚙. La app pasará a modo «Solo local».')
add_spacer(doc)

add_heading(doc, '6.4 Copias de seguridad', 2)
add_bullet(doc, 'Google Sheets (celda A1) es la copia compartida. Descargable en cualquier momento como CSV/XLSX.')
add_bullet(doc, 'Cada navegador tiene su propia copia en localStorage — no confundir con una copia de seguridad real.')
add_bullet(doc, 'Para exportar manualmente: DevTools → Application → Local Storage → copiar valor de coursetrack_v1.')
add_spacer(doc)

add_heading(doc, '6.5 Flujo de trabajo recomendado por usuario', 2)
add_para(doc, 'Al abrir la aplicación por primera vez en un navegador:')
add_bullet(doc, 'Seleccionar el nombre de usuario en el modal inicial.', level=1)
add_bullet(doc, 'Ir a ⚙ Configuración e introducir la URL del Apps Script.', level=1)
add_bullet(doc, 'Pulsar Guardar y sincronizar — la app descargará los datos compartidos.', level=1)
add_spacer(doc, 4)
add_para(doc, 'En cada sesión de trabajo:')
add_bullet(doc, 'La app carga datos locales inmediatamente y luego sincroniza con Sheets (indicador 🟢 en cabecera).', level=1)
add_bullet(doc, 'Si el indicador muestra ⚫ Local, los cambios no se compartirán con el equipo.', level=1)
add_bullet(doc, 'Si muestra 🔴 Sin conexión, los cambios se guardan en local. Recargar para sincronizar después.', level=1)
add_spacer(doc)

add_heading(doc, '6.6 Resolución de problemas frecuentes', 2)
add_spacer(doc, 4)
make_table(doc,
    ['Síntoma', 'Causa probable', 'Solución'],
    [
        ['No aparece el selector de usuario',
         'Ya hay un usuario en localStorage',
         'Borrar coursetrack_user en DevTools → Application → Local Storage'],
        ['App muestra ⚫ Local aunque hay URL configurada',
         'URL del GAS incorrecta o script no desplegado',
         'Revisar URL en ⚙; redesplegar el Apps Script si es necesario'],
        ['Cambios de otro usuario no aparecen',
         'Copia local más reciente (last-write-wins)',
         'Recargar la página para forzar la descarga desde Sheets'],
        ['Datos distintos en dos navegadores',
         'Uno no tiene la URL de GAS configurada',
         'Configurar la URL en el navegador sin conexión y sincronizar'],
        ['Indicador muestra 🔴 Sin conexión',
         'Error de red o CORS bloqueado',
         'Comprobar conectividad; verificar que el GAS está publicado como «Cualquiera»'],
    ],
    [5, 5, 6],
)
add_spacer(doc, 12)

# ── 7. Guía de despliegue ─────────────────────────────────────────────────
add_heading(doc, '7. Guía de despliegue', 1)

add_heading(doc, '7.1 Publicar cambios en GitHub Pages', 2)
add_code_block(doc, [
    '# Desde la carpeta del repositorio local',
    'git add index.html',
    'git commit -m "descripción del cambio"',
    'git push',
    '',
    '# URL pública: https://pacosirvent80-hash.github.io/accenture-coursetrack/',
])

add_heading(doc, '7.2 Configurar Google Apps Script (primera vez)', 2)
add_para(doc, 'El código del Apps Script se obtiene desde el modal ⚙ de la propia aplicación.')
add_bullet(doc, 'Crear una Google Sheet con dos pestañas: Data y Log.')
add_bullet(doc, 'En la hoja: Extensiones → Apps Script → pegar el código del modal.')
add_bullet(doc, 'Implementar → Nueva implementación → Aplicación web.')
add_bullet(doc, 'Ejecutar como: Yo.', level=1)
add_bullet(doc, 'Acceso: Cualquiera (sin iniciar sesión).', level=1)
add_bullet(doc, 'Copiar la URL del Web App generada.')
add_bullet(doc, 'En la app: ⚙ → pegar URL → Guardar y sincronizar.')
add_spacer(doc)

add_heading(doc, '7.3 Actualizar la guía de usuario Word', 2)
add_code_block(doc, [
    'cd C:\\...\\accenture-coursetrack',
    'python gen_guia_coursetrack.py   # genera CourseTrack_Guia_Usuario.docx',
    'python gen_coursetrack_tech.py   # genera CourseTrack_EspecificacionesTecnicas.docx',
])

# ── 8. Referencia rápida ──────────────────────────────────────────────────
add_heading(doc, '8. Referencia rápida', 1)
add_spacer(doc, 4)
make_table(doc,
    ['Elemento', 'Valor / Ruta'],
    [
        ['Repositorio GitHub',     'github.com/pacosirvent80-hash/accenture-coursetrack'],
        ['URL pública',            'pacosirvent80-hash.github.io/accenture-coursetrack/'],
        ['Fichero principal',      'index.html (≈ 85 KB, HTML + CSS + JS)'],
        ['Guía de usuario',        'CourseTrack_Guia_Usuario.docx'],
        ['Guía técnica',           'CourseTrack_EspecificacionesTecnicas.docx'],
        ['Script guía usuario',    'gen_guia_coursetrack.py'],
        ['Script guía técnica',    'gen_coursetrack_tech.py'],
        ['localStorage — datos',   'coursetrack_v1'],
        ['localStorage — GAS URL', 'coursetrack_gas_url'],
        ['localStorage — usuario', 'coursetrack_user'],
        ['Arrays de config.',      'index.html, ≈ línea 806: USERS, PROVIDERS, LANGUAGES'],
        ['Rama de publicación',    'main (raíz /)'],
    ],
    [6, 10],
)

# ── Guardar ───────────────────────────────────────────────────────────────
output = 'CourseTrack_EspecificacionesTecnicas.docx'
doc.save(output)
print(f'✓ Documento generado: {output}')
