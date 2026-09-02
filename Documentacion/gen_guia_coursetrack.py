# -*- coding: utf-8 -*-
"""Genera CourseTrack_Guia_Usuario.docx en la misma carpeta."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Paleta de colores ──────────────────────────────────────────────────────────
AZUL_MARCA   = RGBColor(0x3B, 0x82, 0xF6)   # primary azul CourseTrack
AZUL_OSCURO  = RGBColor(0x1E, 0x29, 0x3B)   # texto principal
GRIS_TEXTO   = RGBColor(0x64, 0x74, 0x8B)   # texto secundario
ROJO         = RGBColor(0xEF, 0x44, 0x44)   # incidencia
VERDE        = RGBColor(0x22, 0xC5, 0x5E)   # aprobado
NARANJA      = RGBColor(0xF9, 0x73, 0x16)   # en corrección
MORADO       = RGBColor(0x8B, 0x5C, 0xF6)   # en revisión
TEAL         = RGBColor(0x0D, 0x94, 0x88)   # publicado
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_para(doc, text='', bold_parts=None, color=None, space_after=6):
    """Añade párrafo con texto. bold_parts = lista de strings que irán en negrita."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if not bold_parts:
        run = p.add_run(text)
        run.font.size  = Pt(10.5)
        if color:
            run.font.color.rgb = color
    else:
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx == -1:
                continue
            before = remaining[:idx]
            after  = remaining[idx + len(bp):]
            if before:
                r = p.add_run(before)
                r.font.size = Pt(10.5)
            r = p.add_run(bp)
            r.bold = True
            r.font.size = Pt(10.5)
            remaining = after
        if remaining:
            r = p.add_run(remaining)
            r.font.size = Pt(10.5)
    return p


def add_bullet(doc, text, level=0, bold_start=None):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    try:
        p = doc.add_paragraph(style=style)
    except Exception:
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        p.paragraph_format.left_indent = Cm(0.8 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    if bold_start:
        r = p.add_run(bold_start)
        r.bold = True
        r.font.size = Pt(10.5)
        rest = text[len(bold_start):]
        if rest:
            r2 = p.add_run(rest)
            r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
    return p


def add_note(doc, text, tipo='info'):
    """Caja de nota con fondo coloreado."""
    colors = {'info': 'EFF6FF', 'warn': 'FEF9C3', 'danger': 'FEF2F2'}
    bg = colors.get(tipo, 'EFF6FF')
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(10)
    doc.add_paragraph()   # espacio tras la tabla
    return tbl


def add_status_table(doc):
    """Tabla de estados con colores."""
    estados = [
        ('📥 Entregado',     'Curso recibido del proveedor. Pendiente de revisión.',                              '6366F1'),
        ('🔍 En revisión',   'La oficina técnica está validando el contenido y el funcionamiento técnico.',        '8B5CF6'),
        ('⚡ Incidencia',    'Se han detectado errores. El curso ha sido devuelto al proveedor.',                  'EF4444'),
        ('🔧 En corrección', 'El proveedor está aplicando las correcciones indicadas en el informe de incidencia.','F97316'),
        ('✅ Aprobado',      'El curso ha superado la validación. Listo para publicar en el LMS.',                 '22C55E'),
        ('🚀 Publicado',     'Curso visible y activo en la plataforma de formación.',                              '0D9488'),
        ('⏸ En pausa',      'El proceso está temporalmente detenido (espera legal, revisión de contenido, etc.)', '94A3B8'),
    ]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Estado', 'Descripción', 'Color']):
        set_cell_bg(hdr[i], '1E293B')
        hdr[i].paragraphs[0].clear()
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = BLANCO
        r.font.size = Pt(10)

    for estado, desc, color in estados:
        row = tbl.add_row().cells
        set_cell_bg(row[0], color + '22')   # fondo muy suave
        row[0].paragraphs[0].add_run(estado).font.size = Pt(10)
        row[1].paragraphs[0].add_run(desc).font.size   = Pt(10)
        set_cell_bg(row[2], color)
        row[2].paragraphs[0].add_run('').font.size      = Pt(10)

    doc.add_paragraph()
    return tbl


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Fuente base
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

# ── PORTADA ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('📋  CourseTrack')
r.font.size  = Pt(28)
r.bold       = True
r.font.color.rgb = AZUL_MARCA

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Gestor de Cursos eLearning')
r2.font.size  = Pt(16)
r2.font.color.rgb = GRIS_TEXTO

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Guía de usuario')
r3.font.size = Pt(13)
r3.italic    = True

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run('Accenture · Oficina Técnica de Formación eLearning')
r4.font.size = Pt(11)
r4.font.color.rgb = GRIS_TEXTO

doc.add_page_break()

# ── 1. INTRODUCCIÓN ───────────────────────────────────────────────────────────
add_heading(doc, '1. ¿Qué es CourseTrack?', 1, AZUL_MARCA)
add_para(doc,
    'CourseTrack es una aplicación web diseñada para la oficina técnica de formación. '
    'Su objetivo es dar visibilidad completa al ciclo de vida de cada curso eLearning: '
    'desde que el proveedor hace la primera entrega hasta que el curso queda publicado '
    'en el LMS, pasando por todas las iteraciones de revisión, corrección y aprobación.',
    space_after=6)
add_para(doc,
    'La herramienta está concebida para ser compartida por un equipo pequeño sin necesidad '
    'de servidores propios ni bases de datos. Un único fichero HTML alojado en GitHub Pages '
    'actúa como frontend; Google Sheets actúa como backend compartido a través de '
    'Google Apps Script.',
    space_after=6)

add_heading(doc, '2. Acceso', 1, AZUL_MARCA)
tbl = doc.add_table(rows=4, cols=2)
tbl.style = 'Table Grid'
data = [
    ('Aplicación web',         'https://pacosirvent80-hash.github.io/accenture-coursetrack/'),
    ('Repositorio GitHub',     'https://github.com/pacosirvent80-hash/accenture-coursetrack'),
    ('Google Apps Script URL', 'https://script.google.com/macros/s/AKfycbziQw_JwEW5MZDlOk3\nAaBLzTfhtfB3U4riXq_NnU1Fzy2VTJTITllKWNlAi7Jwvy2d0/exec'),
    ('Entorno de pruebas SCORM', 'SCORM Cloud (https://cloud.scorm.com)'),
]
set_cell_bg(tbl.rows[0].cells[0], '1E293B')
set_cell_bg(tbl.rows[0].cells[1], '1E293B')
for i, (k, v) in enumerate(data):
    row = tbl.rows[i].cells
    if i == 0:
        rk = row[0].paragraphs[0].add_run(k)
        rk.bold = True; rk.font.color.rgb = BLANCO; rk.font.size = Pt(10)
        rv = row[1].paragraphs[0].add_run(v)
        rv.bold = True; rv.font.color.rgb = BLANCO; rv.font.size = Pt(10)
    else:
        row[0].paragraphs[0].add_run(k).font.size = Pt(10)
        rv = row[1].paragraphs[0].add_run(v)
        rv.font.size = Pt(9.5)
        rv.font.color.rgb = AZUL_MARCA

doc.add_paragraph()
add_note(doc,
    '⚠  La URL del Apps Script otorga acceso de lectura y escritura sin autenticación. '
    'No incluirla en correos masivos ni en canales públicos. Cada usuario la introduce '
    'manualmente en la configuración de su navegador.', 'warn')

# ── 3. INTERFAZ PRINCIPAL ─────────────────────────────────────────────────────
add_heading(doc, '3. Interfaz principal', 1, AZUL_MARCA)
add_para(doc,
    'Al abrir la aplicación se muestra la pantalla principal, dividida en cuatro zonas:')

add_heading(doc, '3.1 Cabecera', 2)
add_para(doc,
    'Barra fija en la parte superior. Contiene el logotipo, el indicador de conexión y '
    'los botones de acción global.')
add_bullet(doc, 'Indicador de conexión: muestra si la app está en modo Local (sin Google Sheets), '
           'sincronizando, conectada (🟢) o con error (🔴). Hacer clic abre la configuración.',
           bold_start='Indicador de conexión:')
add_bullet(doc, '⬇ Exportar: descarga todos los datos actuales en un fichero JSON.', bold_start='⬇ Exportar:')
add_bullet(doc, '⬆ Importar: carga datos desde un fichero JSON previamente exportado.', bold_start='⬆ Importar:')
add_bullet(doc, '⚙: abre el panel de configuración de Google Sheets.', bold_start='⚙:')
add_bullet(doc, '+ Nuevo curso: abre el formulario de alta de un nuevo curso.', bold_start='+ Nuevo curso:')

add_heading(doc, '3.2 Panel de estadísticas', 2)
add_para(doc,
    'Seis tarjetas numéricas que muestran en tiempo real cuántos cursos hay en cada estado. '
    'Hacer clic en cualquier tarjeta filtra la lista de cursos por ese estado.')
add_bullet(doc, 'Total — número total de cursos registrados.', bold_start='Total')
add_bullet(doc, 'En revisión — cursos que la oficina técnica está validando actualmente.', bold_start='En revisión')
add_bullet(doc, 'Incidencias — cursos devueltos al proveedor con errores abiertos.', bold_start='Incidencias')
add_bullet(doc, 'En corrección — cursos en los que el proveedor está aplicando correcciones.', bold_start='En corrección')
add_bullet(doc, 'Aprobados — cursos validados y listos para publicar.', bold_start='Aprobados')
add_bullet(doc, 'Publicados — cursos activos en el LMS.', bold_start='Publicados')

add_heading(doc, '3.3 Barra de filtros', 2)
add_para(doc,
    'Situada bajo las estadísticas, permite acotar qué cursos se ven en la lista:')
add_bullet(doc, 'Buscador de texto: filtra por nombre, código o proveedor (búsqueda en tiempo real).', bold_start='Buscador de texto:')
add_bullet(doc, 'Filtro de estado: desplegable con todos los estados posibles.', bold_start='Filtro de estado:')
add_bullet(doc, 'Filtro de proveedor: desplegable que se construye automáticamente con los proveedores registrados.', bold_start='Filtro de proveedor:')
add_bullet(doc, 'Contador: muestra cuántos cursos se muestran del total.', bold_start='Contador:')

add_heading(doc, '3.4 Lista de cursos', 2)
add_para(doc,
    'En escritorio se muestra como tabla con columnas: Curso (nombre + código), Proveedor, '
    'Versión, Estado, Incidencias abiertas, Días en el estado actual y Última actualización. '
    'En móvil cada curso se muestra como una tarjeta con nombre, estado e incidencias. '
    'Las columnas Proveedor, Versión y Días se ocultan automáticamente en pantallas pequeñas. '
    'Hacer clic en un curso abre el panel de detalle.')

# ── 4. GESTIÓN DE CURSOS ──────────────────────────────────────────────────────
add_heading(doc, '4. Gestión de cursos', 1, AZUL_MARCA)

add_heading(doc, '4.1 Crear un curso', 2)
add_para(doc,
    'Pulsar el botón + Nuevo curso en la cabecera abre un formulario modal con los campos:')
add_bullet(doc, 'Nombre del curso (obligatorio)')
add_bullet(doc, 'Código interno (obligatorio, p. ej. CIB-001)')
add_bullet(doc, 'Proveedor (obligatorio) — selector con la lista predefinida: Accenture, INAP, Mercadona, Test Provider.')
add_bullet(doc, 'Estándar técnico: SCORM 1.2, SCORM 2004, xAPI (Tin Can), cmi5 o HTML sin tracking')
add_bullet(doc, 'Duración estimada (p. ej. 45 min)')
add_bullet(doc, 'Idioma — selector con Español, Inglés, Catalán, Euskera y Gallego (por defecto Español).')
add_bullet(doc, 'Estado inicial (seleccionable entre los siete estados del flujo)')
add_bullet(doc, 'Notas internas (observaciones para la oficina técnica)')
add_para(doc,
    'Al confirmar, el curso se registra con la versión 1.0 y el estado elegido. '
    'Los datos se guardan inmediatamente en localStorage y, si hay conexión configurada, '
    'también en Google Sheets.',
    space_after=8)

add_heading(doc, '4.2 Editar un curso', 2)
add_para(doc,
    'Para modificar los datos de un curso ya creado, abrir su panel de detalle y pulsar '
    '✏ Editar curso en la barra de acciones inferior. El modal se abre con todos los '
    'campos pre-rellenados y permite modificar:')
add_bullet(doc, 'Nombre del curso')
add_bullet(doc, 'Código interno')
add_bullet(doc, 'Proveedor (selector predefinido)')
add_bullet(doc, 'Estándar SCORM')
add_bullet(doc, 'Duración estimada')
add_bullet(doc, 'Idioma (selector predefinido)')
add_bullet(doc, 'Notas internas')
add_para(doc,
    'Al confirmar, los cambios se guardan inmediatamente y se sincronizan con Google Sheets.',
    space_after=8)

add_heading(doc, '4.3 Eliminar un curso', 2)
add_para(doc,
    'Desde el panel de detalle del curso, el botón 🗑 (esquina inferior derecha) elimina el curso '
    'tras pedir confirmación. Esta acción es irreversible en la sesión actual; si se trabaja '
    'con Google Sheets, los datos se sobreescriben en la próxima sincronización.',
    space_after=8)

# ── 5. ESTADOS ────────────────────────────────────────────────────────────────
add_heading(doc, '5. Ciclo de vida y estados', 1, AZUL_MARCA)
add_para(doc,
    'Cada curso pasa por estados que reflejan su situación en el flujo de validación. '
    'Al cambiar de estado se crea automáticamente una nueva versión menor del curso '
    '(p. ej. de v1.0 a v1.1), manteniendo el historial completo.')
doc.add_paragraph()
add_status_table(doc)
add_para(doc,
    'El flujo habitual es: Entregado → En revisión → (si hay errores) Incidencia → '
    'En corrección → En revisión → Aprobado → Publicado. '
    'En pausa puede aplicarse en cualquier momento del ciclo.',
    color=GRIS_TEXTO, space_after=8)

# ── 6. PANEL DE DETALLE ───────────────────────────────────────────────────────
add_heading(doc, '6. Panel de detalle del curso', 1, AZUL_MARCA)
add_para(doc,
    'Al seleccionar un curso se abre el panel lateral (en móvil ocupa toda la pantalla). '
    'Contiene tres pestañas:')

add_heading(doc, '6.1 Pestaña Info', 2)
add_para(doc,
    'Muestra los datos del curso: proveedor, estándar técnico, idioma, duración, '
    'fecha de alta, última actualización y notas internas.')

add_heading(doc, '6.2 Pestaña Versiones', 2)
add_para(doc,
    'Historial cronológico inverso de todas las versiones del curso, presentado como '
    'una línea de tiempo. Cada entrada muestra: número de versión, estado, fecha y '
    'notas del cambio. Si esa versión tuvo incidencias se indica el número.')

add_heading(doc, '6.3 Pestaña Incidencias', 2)
add_para(doc,
    'Listado completo de incidencias del curso, divididas en Abiertas y Resueltas. '
    'Cada tarjeta de incidencia muestra:')
add_bullet(doc, 'Categoría (SCORM/xAPI, Contenido, Diseño/UX, Técnico, Estructura IMS, Evaluación, Vídeo/Audio, Accesibilidad, Otro)')
add_bullet(doc, 'Severidad: Alta, Media o Baja')
add_bullet(doc, 'Versión en la que se detectó')
add_bullet(doc, 'Título y descripción detallada')
add_bullet(doc, 'Fecha de detección y, si está resuelta, texto de resolución')
add_para(doc,
    'Cada tarjeta incluye una barra de acciones en la parte inferior con los botones '
    '✏ Editar incidencia (disponible siempre) y Resolver ✓ (solo en incidencias abiertas).',
    space_after=8)

add_heading(doc, '6.4 Acciones desde el panel', 2)
add_para(doc, 'La barra de acciones en la parte inferior del panel contiene:')
add_bullet(doc, '✏ Editar curso: abre el modal de edición con todos los campos del curso pre-rellenados.', bold_start='✏ Editar curso:')
add_bullet(doc, '↔ Cambiar estado: mueve el curso al siguiente estado del ciclo y crea una versión menor.', bold_start='↔ Cambiar estado:')
add_bullet(doc, '+ Nueva versión: registra una reentrega mayor del proveedor (p. ej. de v1.x a v2.0).', bold_start='+ Nueva versión:')
add_bullet(doc, '⚡ Nueva incidencia: abre el formulario para registrar un nuevo error detectado en la revisión.', bold_start='⚡ Nueva incidencia:')
add_bullet(doc, '🗑: elimina el curso completo (con confirmación).', bold_start='🗑:')

# ── 7. REGISTRAR INCIDENCIA ───────────────────────────────────────────────────
add_heading(doc, '7. Gestión de incidencias', 1, AZUL_MARCA)
add_para(doc,
    'Las incidencias registran los errores o problemas detectados durante la revisión '
    'de un curso. Cada incidencia queda asociada a la versión activa del curso en el '
    'momento del registro y se muestra en la pestaña Incidencias del panel de detalle.')

add_heading(doc, '7.1 Registrar una nueva incidencia', 2)
add_para(doc,
    'Desde el panel de detalle → botón ⚡ Nueva incidencia. El formulario solicita:')
add_bullet(doc, 'Categoría: tipo de problema detectado (técnico, contenido, accesibilidad, etc.)')
add_bullet(doc, 'Severidad: Alta (bloquea la publicación), Media (requiere corrección antes de publicar) o Baja (mejora no urgente).')
add_bullet(doc, 'Título: resumen en una línea del problema.')
add_bullet(doc, 'Descripción detallada: pasos para reproducir, entorno afectado, impacto en el alumno.')
add_bullet(doc, 'Fecha de detección.')
add_para(doc,
    'Al confirmar, la app cambia automáticamente a la pestaña Incidencias del panel.',
    space_after=8)

add_heading(doc, '7.2 Editar una incidencia', 2)
add_para(doc,
    'Todas las incidencias (abiertas y resueltas) pueden editarse en cualquier momento. '
    'En la pestaña Incidencias, cada tarjeta muestra el botón ✏ Editar incidencia '
    'en su barra de acciones inferior. El modal de edición incluye todos los campos:')
add_bullet(doc, 'Categoría y severidad')
add_bullet(doc, 'Título y descripción detallada')
add_bullet(doc, 'Fecha de detección')
add_bullet(doc, 'Estado (Abierta / Resuelta) — al marcar como Resuelta aparece el campo de texto de resolución')
add_bullet(doc, 'Texto de resolución (editable si la incidencia está resuelta)')
add_para(doc,
    'Los campos de trazabilidad (Creada por, Última edición, Resuelta por) se muestran '
    'en modo solo lectura dentro del modal. Al guardar, se registran editadoPor '
    'y fechaEdicion con el usuario activo y la fecha actual.',
    color=GRIS_TEXTO, space_after=8)

add_heading(doc, '7.3 Resolver una incidencia', 2)
add_para(doc,
    'En la pestaña Incidencias, cada tarjeta abierta muestra el botón Resolver ✓ '
    'en su barra de acciones. '
    'Al pulsarlo se pide un texto de resolución (opcional) y la incidencia pasa a Resuelta, '
    'con su tarjeta marcada en verde. Las incidencias resueltas se conservan en el historial; '
    'no se eliminan.',
    space_after=8)

# ── 8. CONTROL DE VERSIONES ───────────────────────────────────────────────────
add_heading(doc, '8. Control de versiones', 1, AZUL_MARCA)
add_para(doc,
    'CourseTrack distingue dos tipos de versión:')
add_bullet(doc,
    'Versión menor (automática): se crea cada vez que se cambia el estado del curso. '
    'El número sube en el segundo dígito (1.0 → 1.1 → 1.2...).',
    bold_start='Versión menor (automática):')
add_bullet(doc,
    'Versión mayor (manual): se crea con el botón + Nueva versión para registrar '
    'una reentrega completa del proveedor. El primer dígito sube y el segundo se reinicia '
    '(1.x → 2.0).',
    bold_start='Versión mayor (manual):')
add_para(doc,
    'Cada versión tiene su propio estado, fecha, notas y lista de incidencias. '
    'El historial completo es inmutable — no se pueden editar versiones pasadas.',
    color=GRIS_TEXTO, space_after=8)

# ── 9. IDENTIFICACIÓN DE USUARIO ─────────────────────────────────────────────
add_heading(doc, '9. Identificación de usuario y trazabilidad', 1, AZUL_MARCA)
add_para(doc,
    'CourseTrack incluye un sistema de identificación ligero pensado para equipos pequeños. '
    'No requiere contraseñas ni servidor de autenticación: el usuario elige su nombre de una '
    'lista predefinida y ese nombre queda registrado en cada acción que realice.')

add_heading(doc, '9.1 Selección de usuario al acceder', 2)
add_para(doc,
    'La primera vez que se abre la aplicación en un navegador (o cuando se borra el '
    'localStorage), aparece un modal de selección de usuario que no puede cerrarse sin '
    'elegir un nombre. Una vez elegido, queda guardado en el navegador y no vuelve a '
    'pedirse en siguientes visitas.')
add_bullet(doc,
    'El chip 👤 de la cabecera muestra el nombre activo y permite cambiarlo en cualquier momento.',
    bold_start='El chip 👤 de la cabecera')
add_bullet(doc,
    'Cada navegador/dispositivo recuerda su propio usuario de forma independiente.',
    bold_start='Cada navegador/dispositivo')

add_heading(doc, '9.2 Campos de trazabilidad', 2)
add_para(doc, 'Cada acción registrada incluye el nombre del usuario activo:')
tbl_usr = doc.add_table(rows=4, cols=2)
tbl_usr.style = 'Table Grid'
set_cell_bg(tbl_usr.rows[0].cells[0], '1E293B')
set_cell_bg(tbl_usr.rows[0].cells[1], '1E293B')
usr_data = [
    ('Acción', 'Campo guardado'),
    ('Crear curso / cambiar estado / registrar nueva versión', 'modificadoPor  (en el objeto de versión)'),
    ('Registrar incidencia', 'creadaPor  (en el objeto de incidencia)'),
    ('Resolver incidencia', 'resueltaPor  (en el objeto de incidencia)'),
]
for i, (k, v) in enumerate(usr_data):
    row = tbl_usr.rows[i].cells
    if i == 0:
        for ci, txt in [(0, k), (1, v)]:
            r = row[ci].paragraphs[0].add_run(txt)
            r.bold = True; r.font.color.rgb = BLANCO; r.font.size = Pt(10)
    else:
        row[0].paragraphs[0].add_run(k).font.size = Pt(10)
        row[1].paragraphs[0].add_run(v).font.size  = Pt(10)
doc.add_paragraph()
add_para(doc,
    'Estos campos forman parte del JSON que se sincroniza con Google Sheets, '
    'por lo que la trazabilidad queda grabada de forma permanente en la hoja de cálculo.',
    color=GRIS_TEXTO, space_after=8)

add_heading(doc, '9.3 Filtrar por usuario', 2)
add_para(doc,
    'El toolbar principal incluye un desplegable "Todos los usuarios" que se autocompleta '
    'automáticamente con los nombres que han realizado al menos un cambio de estado. '
    'Permite ver de un vistazo qué cursos ha tocado cada miembro del equipo.')
add_bullet(doc, 'La columna "Por" de la tabla desktop muestra el responsable del último cambio de estado.')
add_bullet(doc, 'En móvil, la tarjeta de cada curso muestra el primer nombre del último modificador.')
add_bullet(doc, 'Todos los encabezados de columna son ordenables con un clic.')

add_heading(doc, '9.4 Gestión de las listas predefinidas', 2)
add_para(doc,
    'Tres arrays hardcodeados en index.html (~línea 792) controlan los valores disponibles '
    'en los selectores de la aplicación:')
add_bullet(doc, 'USERS — miembros del equipo que pueden identificarse al acceder.', bold_start='USERS')
add_bullet(doc, 'PROVIDERS — proveedores disponibles en el filtro del toolbar y en el formulario de alta de curso.', bold_start='PROVIDERS')
add_bullet(doc, 'LANGUAGES — idiomas disponibles en el formulario de alta de curso (el primero es el valor por defecto).', bold_start='LANGUAGES')
add_para(doc, 'Para añadir o eliminar valores en cualquiera de las tres listas:')
add_bullet(doc, 'Abrir index.html y localizar los arrays (aproximadamente en la línea 792).')
add_bullet(doc, 'Editar la lista de valores.')
add_bullet(doc, 'Hacer git push. GitHub Pages actualiza en 1–2 minutos.')
add_note(doc,
    '⚠  La selección de usuario no es autenticación: cualquier persona con acceso a la URL '
    'puede elegir cualquier nombre de la lista. Está diseñado para equipos de confianza '
    'donde la trazabilidad importa más que la seguridad de acceso.',
    'warn')

# ── 10. GOOGLE SHEETS ─────────────────────────────────────────────────────────
add_heading(doc, '10. Sincronización con Google Sheets', 1, AZUL_MARCA)

add_heading(doc, '10.1 Arquitectura de datos', 2)
add_para(doc,
    'La aplicación usa Google Apps Script como API REST mínima. La hoja de cálculo '
    'tiene dos pestañas:')
add_bullet(doc,
    'Data (celda A1): contiene todos los cursos serializados en un único JSON. '
    'Ejemplo: [{"id":"abc","name":"Curso X","versions":[...]},...]',
    bold_start='Data (celda A1):')
add_bullet(doc,
    'Log: cada vez que la app guarda datos, se añade una fila con timestamp y tamaño del JSON en bytes.',
    bold_start='Log:')
add_note(doc,
    'ℹ  El modelo es last-write-wins: si dos personas guardan al mismo tiempo, prevalece la última escritura. '
    'Para equipos pequeños (hasta 5 personas) esto es suficiente. El Log permite detectar escrituras simultáneas.',
    'info')

add_heading(doc, '10.2 Configurar el Apps Script paso a paso', 2)
add_para(doc, 'Paso 1 — Preparar la hoja de cálculo')
add_bullet(doc, 'Abrir Google Sheets y crear una hoja nueva (o usar una existente).')
add_bullet(doc, 'Renombrar la primera pestaña como Data y crear una segunda pestaña llamada Log.')

add_para(doc, 'Paso 2 — Crear el Apps Script')
add_bullet(doc, 'En la hoja: menú Extensiones → Apps Script.')
add_bullet(doc, 'Borrar el contenido del editor y pegar el código que aparece en ⚙ Configuración → (sección "Código del Apps Script").')
add_bullet(doc, 'Guardar el proyecto (nombre sugerido: CourseTrack API).')

add_para(doc, 'Paso 3 — Desplegar como Web App')
add_bullet(doc, 'Implementar → Nueva implementación.')
add_bullet(doc, 'Tipo: Aplicación web.')
add_bullet(doc, 'Ejecutar como: Yo (tu cuenta de Google).')
add_bullet(doc, 'Acceso: Cualquiera (sin iniciar sesión). Este punto es crítico — sin él la app no puede leer/escribir.')
add_bullet(doc, 'Pulsar Implementar y autorizar los permisos solicitados.')
add_bullet(doc, 'Copiar la URL del Web App generada.')

add_para(doc, 'Paso 4 — Conectar CourseTrack')
add_bullet(doc, 'Abrir la aplicación web en el navegador.')
add_bullet(doc, 'Pulsar el indicador de conexión (esquina derecha de la cabecera) o el icono ⚙.')
add_bullet(doc, 'Pegar la URL del Web App en el campo "URL del Web App".')
add_bullet(doc, 'Pulsar Probar para verificar la conexión.')
add_bullet(doc, 'Pulsar Guardar y sincronizar. El indicador cambia a 🟢 Google Sheets.')
add_note(doc,
    '⚠  Esta configuración se guarda solo en el navegador actual (localStorage). '
    'Cada persona del equipo debe realizar el paso 4 en su propio navegador.',
    'warn')

add_heading(doc, '10.3 Comportamiento del indicador de conexión', 2)
tbl2 = doc.add_table(rows=5, cols=2)
tbl2.style = 'Table Grid'
set_cell_bg(tbl2.rows[0].cells[0], '1E293B')
set_cell_bg(tbl2.rows[0].cells[1], '1E293B')
hdata = [('Estado', 'Significado')]
cdata = [
    ('⚫ Local',          'No hay URL de Apps Script configurada. Los datos se guardan solo en este navegador.'),
    ('⟳ Sincronizando',  'Se está enviando o recibiendo información de Google Sheets.'),
    ('🟢 Google Sheets', 'Última operación de sincronización completada con éxito.'),
    ('🔴 Sin conexión',  'Error al conectar con Google Sheets. Los datos se conservan en localStorage como respaldo.'),
]
for i, (k, v) in enumerate(hdata + cdata):
    row = tbl2.rows[i].cells
    if i == 0:
        for ci, txt in [(0, k), (1, v)]:
            r = row[ci].paragraphs[0].add_run(txt)
            r.bold = True; r.font.color.rgb = BLANCO; r.font.size = Pt(10)
    else:
        row[0].paragraphs[0].add_run(k).font.size = Pt(10)
        row[1].paragraphs[0].add_run(v).font.size = Pt(10)
doc.add_paragraph()

add_heading(doc, '10.4 Flujo de guardado', 2)
add_para(doc,
    'Cada vez que se modifica algún dato (nuevo curso, cambio de estado, incidencia, resolución):')
add_bullet(doc, 'Se guarda inmediatamente en localStorage del navegador actual.')
add_bullet(doc, 'Si hay URL de Apps Script configurada, se envía también a Google Sheets (POST asíncrono).')
add_para(doc,
    'Al cargar la aplicación: se muestran primero los datos locales (carga instantánea) y '
    'después se intenta cargar desde Google Sheets. Si el remoto tiene datos más recientes, '
    'los locales se actualizan.',
    color=GRIS_TEXTO, space_after=8)

# ── 10. IMPORTAR / EXPORTAR ───────────────────────────────────────────────────
add_heading(doc, '11. Importar y exportar datos', 1, AZUL_MARCA)
add_bullet(doc,
    '⬇ Exportar: descarga un fichero JSON con todos los cursos. Útil como copia de seguridad '
    'o para migrar datos entre navegadores.',
    bold_start='⬇ Exportar:')
add_bullet(doc,
    '⬆ Importar: carga un JSON previamente exportado, reemplazando todos los datos actuales. '
    'Pide confirmación antes de sobreescribir.',
    bold_start='⬆ Importar:')
add_note(doc,
    'ℹ  Si el equipo trabaja con Google Sheets como fuente de verdad compartida, '
    'no es necesario exportar/importar manualmente — la sincronización es automática. '
    'Importar es útil para restaurar un estado anterior desde una copia de seguridad.',
    'info')

# ── 11. MÓVIL ─────────────────────────────────────────────────────────────────
add_heading(doc, '12. Uso desde dispositivos móviles', 1, AZUL_MARCA)
add_para(doc,
    'La aplicación es completamente funcional en smartphones y tablets Android e iOS. '
    'El diseño se adapta automáticamente:')
add_bullet(doc, 'Cabecera: se muestran solo los iconos de los botones (sin texto) para maximizar el espacio.')
add_bullet(doc, 'Lista de cursos: cada curso se muestra como tarjeta con nombre, estado e incidencias abiertas.')
add_bullet(doc, 'Filtros: búsqueda y desplegables apilados verticalmente.')
add_bullet(doc, 'Panel de detalle: ocupa la pantalla completa al seleccionar un curso.')
add_bullet(doc, 'Estadísticas: rejilla de 3 columnas (o 2 en móviles muy pequeños).')
add_para(doc,
    'La URL pública se puede compartir directamente con cualquier miembro del equipo para '
    'consultar el estado de los cursos desde el móvil.',
    space_after=8)

# ── 12. TECNOLOGÍA ────────────────────────────────────────────────────────────
add_heading(doc, '13. Arquitectura técnica', 1, AZUL_MARCA)
tbl3 = doc.add_table(rows=6, cols=2)
tbl3.style = 'Table Grid'
set_cell_bg(tbl3.rows[0].cells[0], '1E293B')
set_cell_bg(tbl3.rows[0].cells[1], '1E293B')
techdata = [
    ('Capa', 'Tecnología'),
    ('Frontend', 'HTML5, CSS3, JavaScript ES2020+ (Vanilla — sin frameworks)'),
    ('Persistencia local', 'localStorage del navegador (claves: coursetrack_v1, coursetrack_user, coursetrack_gas_url)'),
    ('Persistencia compartida', 'Google Sheets + Google Apps Script REST (POST/GET)'),
    ('Hosting', 'GitHub Pages — rama main, raíz /'),
    ('Responsive', 'CSS Grid, Flexbox, media queries + detección JS de viewport'),
]
for i, (k, v) in enumerate(techdata):
    row = tbl3.rows[i].cells
    if i == 0:
        for ci, txt in [(0, k), (1, v)]:
            r = row[ci].paragraphs[0].add_run(txt)
            r.bold = True; r.font.color.rgb = BLANCO; r.font.size = Pt(10)
    else:
        row[0].paragraphs[0].add_run(k).font.size = Pt(10)
        row[1].paragraphs[0].add_run(v).font.size = Pt(10)
doc.add_paragraph()

add_para(doc,
    'No hay servidor de aplicaciones, ni base de datos, ni proceso de build. '
    'El único fichero es index.html. Para actualizar la aplicación basta con editar ese '
    'fichero y hacer git push al repositorio — GitHub Pages lo publica automáticamente.',
    color=GRIS_TEXTO, space_after=8)

# ── 13. ACTUALIZAR LA APLICACIÓN ──────────────────────────────────────────────
add_heading(doc, '14. Actualizar la aplicación', 1, AZUL_MARCA)
add_para(doc,
    'Cualquier cambio en la aplicación (nuevo campo, corrección de texto, nueva funcionalidad) '
    'sigue este flujo:')
add_bullet(doc, 'Clonar el repositorio si no se ha hecho: git clone https://github.com/pacosirvent80-hash/accenture-coursetrack.git')
add_bullet(doc, 'Editar index.html (directamente o con ayuda de Claude Code).')
add_bullet(doc, 'Publicar: git add index.html → git commit -m "descripción" → git push.')
add_bullet(doc, 'GitHub Pages actualiza la URL pública en 1-2 minutos.')
add_para(doc,
    'Si el navegador muestra la versión antigua tras el push, abrir una pestaña de incógnito '
    'o forzar recarga con Ctrl + Shift + R.',
    color=GRIS_TEXTO, space_after=8)

# ── GUARDAR ───────────────────────────────────────────────────────────────────
output_path = 'CourseTrack_Guia_Usuario.docx'
doc.save(output_path)
print(f'✓ Documento generado: {output_path}')
