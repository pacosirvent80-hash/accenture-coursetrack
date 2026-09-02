"""
Genera CourseTrack_Novedades.docx — resumen de nuevas funcionalidades para usuarios.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

AZUL       = RGBColor(0x1E, 0x3A, 0x5F)
AZUL_MED   = RGBColor(0x25, 0x63, 0xEB)
GRIS_HDR   = RGBColor(0xF1, 0xF5, 0xF9)
GRIS_TEXT  = RGBColor(0x1E, 0x29, 0x3B)
VERDE      = RGBColor(0x15, 0x80, 0x3D)
BLANCO     = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color='D1D5DB'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def no_space(paragraph):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after  = Pt(0)

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Estilos base ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Segoe UI'
style.font.size = Pt(10)
style.font.color.rgb = GRIS_TEXT

# ── Cabecera del documento ────────────────────────────────────────────────
hdr = doc.add_paragraph()
hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(hdr)
run = hdr.add_run('CourseTrack')
run.font.name  = 'Segoe UI'
run.font.size  = Pt(26)
run.font.bold  = True
run.font.color.rgb = AZUL

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(sub)
r = sub.add_run('Novedades de la aplicación')
r.font.name  = 'Segoe UI'
r.font.size  = Pt(13)
r.font.color.rgb = AZUL_MED

fecha = doc.add_paragraph()
fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
fecha.paragraph_format.space_after = Pt(18)
r2 = fecha.add_run(f'Actualización: {datetime.date.today().strftime("%d de %B de %Y")}')
r2.font.name  = 'Segoe UI'
r2.font.size  = Pt(9)
r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

# ── Separador ─────────────────────────────────────────────────────────────
sep = doc.add_paragraph()
no_space(sep)
sep.paragraph_format.space_after = Pt(20)
pPr = sep._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '6')
bottom.set(qn('w:color'), '1E3A5F')
pBdr.append(bottom)
pPr.append(pBdr)

# ── Función para sección numerada ─────────────────────────────────────────
def seccion(numero, titulo, descripcion):
    # Número + título
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    num = p.add_run(f'{numero}.  ')
    num.font.name  = 'Segoe UI'
    num.font.size  = Pt(13)
    num.font.bold  = True
    num.font.color.rgb = AZUL_MED
    ttl = p.add_run(titulo)
    ttl.font.name  = 'Segoe UI'
    ttl.font.size  = Pt(13)
    ttl.font.bold  = True
    ttl.font.color.rgb = AZUL
    # Descripción
    d = doc.add_paragraph(descripcion)
    d.paragraph_format.space_before = Pt(2)
    d.paragraph_format.space_after  = Pt(6)
    d.paragraph_format.left_indent  = Cm(0.8)
    d.runs[0].font.name = 'Segoe UI'
    d.runs[0].font.size = Pt(10)
    d.runs[0].font.color.rgb = GRIS_TEXT

def bullet(texto, negrita_hasta=None):
    """Bullet point con texto. negrita_hasta = nº de caracteres en negrita al inicio."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(1.2)
    if negrita_hasta:
        bold_part = texto[:negrita_hasta]
        rest_part = texto[negrita_hasta:]
        rb = p.add_run(bold_part)
        rb.font.bold = True
        rb.font.name = 'Segoe UI'
        rb.font.size = Pt(10)
        rb.font.color.rgb = GRIS_TEXT
        rr = p.add_run(rest_part)
        rr.font.name = 'Segoe UI'
        rr.font.size = Pt(10)
        rr.font.color.rgb = GRIS_TEXT
    else:
        r = p.add_run(texto)
        r.font.name = 'Segoe UI'
        r.font.size = Pt(10)
        r.font.color.rgb = GRIS_TEXT

def nota(texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.8)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:fill'), 'EFF6FF')
    p._p.get_or_add_pPr().append(shading)
    r = p.add_run('ℹ  ' + texto)
    r.font.name  = 'Segoe UI'
    r.font.size  = Pt(9)
    r.font.italic = True
    r.font.color.rgb = AZUL_MED

def tabla_exportacion():
    tbl = doc.add_table(rows=4, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'

    encabezados = ['Formato', 'Contenido exportado']
    for ci, txt in enumerate(encabezados):
        cell = tbl.rows[0].cells[ci]
        set_cell_bg(cell, '1E3A5F')
        set_cell_border(cell, '1E3A5F')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(txt)
        r.font.bold  = True
        r.font.name  = 'Segoe UI'
        r.font.size  = Pt(9)
        r.font.color.rgb = BLANCO

    filas = [
        ('JSON',
         'Todos los datos de la aplicación en formato estructurado: cursos, versiones e incidencias con todos sus campos, incluyendo trazabilidad de usuarios.'),
        ('CSV',
         'Una fila por incidencia con 25 columnas: datos del curso, de la versión y de la incidencia. Las versiones sin incidencias también aparecen como fila. Compatible con Excel, Numbers y Google Sheets.'),
        ('Excel (.xlsx)',
         'Tres pestañas: Cursos (resumen con estado actual), Versiones (historial completo) e Incidencias (todas las incidencias con trazabilidad). Con colores y formato de tabla.'),
    ]
    bg_alt = ['F8FAFC', 'FFFFFF']
    for ri, (fmt, desc) in enumerate(filas, 1):
        bg = bg_alt[ri % 2]
        row = tbl.rows[ri]
        # Columna formato
        c0 = row.cells[0]
        set_cell_bg(c0, bg)
        set_cell_border(c0)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(fmt)
        r0.font.bold  = True
        r0.font.name  = 'Segoe UI'
        r0.font.size  = Pt(9)
        r0.font.color.rgb = AZUL
        # Columna descripción
        c1 = row.cells[1]
        set_cell_bg(c1, bg)
        set_cell_border(c1)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(desc)
        r1.font.name = 'Segoe UI'
        r1.font.size = Pt(9)
        r1.font.color.rgb = GRIS_TEXT

    # Anchos de columna
    tbl.columns[0].width = Cm(3)
    tbl.columns[1].width = Cm(12)
    for row in tbl.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(12)

    sp = doc.add_paragraph()
    no_space(sp)
    sp.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════════════════
#  CONTENIDO
# ══════════════════════════════════════════════════════════════════════════

intro = doc.add_paragraph(
    'Este documento describe las nuevas funcionalidades incorporadas a CourseTrack '
    'en la última actualización. Todas las mejoras están disponibles sin necesidad '
    'de instalar nada: basta con recargar la aplicación en el navegador.'
)
intro.paragraph_format.space_after = Pt(14)
intro.runs[0].font.name = 'Segoe UI'
intro.runs[0].font.size = Pt(10)

# ── 1. Responsive ─────────────────────────────────────────────────────────
seccion(1, 'Diseño adaptable a cualquier dispositivo',
    'La interfaz se adapta automáticamente a escritorio, tablet y móvil. '
    'En pantallas pequeñas, la lista de cursos pasa de tabla a tarjetas apiladas '
    'que muestran toda la información de forma legible. La cabecera reorganiza '
    'sus controles para no requerir scroll horizontal en ningún tamaño de pantalla.')
nota('No es necesario hacer nada: el diseño se ajusta solo al tamaño de la ventana del navegador.')

# ── 2. Usuarios y trazabilidad ────────────────────────────────────────────
seccion(2, 'Sistema de usuarios y trazabilidad de acciones',
    'Cada miembro del equipo elige su nombre al abrir la aplicación por primera vez. '
    'A partir de ese momento, todas las acciones quedan firmadas automáticamente:')
bullet('Nueva versión o cambio de estado', negrita_hasta=30)
bullet('Registro de una incidencia', negrita_hasta=25)
bullet('Edición de una incidencia existente', negrita_hasta=32)
bullet('Resolución de una incidencia', negrita_hasta=26)
p = doc.add_paragraph(
    'El filtro "Todos los usuarios" en la barra superior permite ver solo '
    'los cursos tocados por una persona concreta. Los nombres disponibles '
    'se configuran en el código fuente de la aplicación.'
)
p.paragraph_format.left_indent = Cm(0.8)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(6)
p.runs[0].font.name = 'Segoe UI'
p.runs[0].font.size = Pt(10)

# ── 3. Editar incidencias ─────────────────────────────────────────────────
seccion(3, 'Edición de incidencias',
    'Cada incidencia registrada tiene ahora el botón ✏ Editar incidencia, que abre '
    'un formulario con todos sus campos. Al guardar, se registra automáticamente '
    'quién la editó y en qué fecha, sin perder la información de quién la creó originalmente.')
nota('El campo "Editada por" y "Fecha de edición" se guardan de forma automática al modificar cualquier campo.')

# ── 4. Editar cursos ──────────────────────────────────────────────────────
seccion(4, 'Edición de cursos',
    'El pie del panel lateral de cada curso incluye el botón ✏ Editar curso. '
    'Permite modificar el nombre, código, proveedor, formato SCORM, idioma, '
    'duración y notas sin necesidad de crear una versión nueva.')

# ── 5. Tema oscuro ────────────────────────────────────────────────────────
seccion(5, 'Tema oscuro',
    'El botón ☀/🌙 situado en la cabecera alterna entre el tema claro y el oscuro. '
    'La preferencia queda guardada en el navegador y se mantiene entre sesiones. '
    'Todos los elementos de la interfaz —modales, formularios, panel lateral, '
    'etiquetas de estado— respetan el tema activo.')

# ── 6. Exportación ────────────────────────────────────────────────────────
seccion(6, 'Exportación de datos',
    'El botón ⬇ Exportar de la cabecera despliega tres opciones para descargar '
    'toda la información de la aplicación:')
doc.add_paragraph().paragraph_format.space_after = Pt(4)
tabla_exportacion()
nota('El archivo descargado lleva la fecha del día en el nombre, por ejemplo: coursetrack_2026-09-02.xlsx')

# ── Pie de documento ──────────────────────────────────────────────────────
doc.add_paragraph().paragraph_format.space_before = Pt(20)
pie_sep = doc.add_paragraph()
no_space(pie_sep)
pPr2 = pie_sep._p.get_or_add_pPr()
pBdr2 = OxmlElement('w:pBdr')
top2 = OxmlElement('w:top')
top2.set(qn('w:val'), 'single')
top2.set(qn('w:sz'), '4')
top2.set(qn('w:color'), 'CBD5E1')
pBdr2.append(top2)
pPr2.append(pBdr2)

pie = doc.add_paragraph()
pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(pie)
rp = pie.add_run('CourseTrack · Accenture · Gestión de cursos eLearning')
rp.font.name  = 'Segoe UI'
rp.font.size  = Pt(8)
rp.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

out = 'CourseTrack_Novedades.docx'
doc.save(out)
print(f'✓ Documento generado: {out}')
